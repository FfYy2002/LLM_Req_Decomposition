# -*- coding: utf-8 -*-
"""
需求抽取自动化脚本（基于 Qwen API）

功能说明：
- 从指定目录读取 PDF / DOCX / DOC 格式的需求文档
- 自动将 .doc 转换为 .docx（需 Microsoft Word）
- 调用 Qwen-Max 模型执行 AR（Atomic Requirement）分解
- 识别“高级需求” → “细节需求”配对
- 输出结果至 Excel 文件（分批次保存，避免内存溢出）

作者：minefan
日期：2025-12-03
依赖：见 requirements.txt
"""

import os
import pandas as pd
from openai import OpenAI
import pdfplumber
from docx import Document
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
import re


# =============================================================================
# 全局配置区（建议通过环境变量或配置文件管理）
# =============================================================================

# 输入文档目录（支持 .pdf、.docx、.doc）
INPUT_DIR = r"D:\req"

# 分批处理配置：每 BATCH_SIZE 个有效文档保存一个 Excel 文件
BATCH_SIZE = 10
# 输出目录（用于存放分批生成的 Excel 文件）
OUTPUT_DIR = "output_batches"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 从环境变量获取 DashScope API Key（安全最佳实践）
# set DASHSCOPE_API_KEY=sk-11dea44cf733433ea60b8efc030e90bf
# echo %DASHSCOPE_API_KEY%
DASHSCOPE_API_KEY = os.getenv("DASHSCOPE_API_KEY")
if not DASHSCOPE_API_KEY:
    raise EnvironmentError(
        "❌ 环境变量 DASHSCOPE_API_KEY 未设置。\n"
        "请在运行前执行：\n"
        "  Windows (PowerShell): $env:DASHSCOPE_API_KEY='sk-xxx'\n"
        "  Linux/macOS: export DASHSCOPE_API_KEY='sk-xxx'"
    )

# 初始化 OpenAI 兼容客户端（指向阿里云 DashScope）
CLIENT = OpenAI(
    api_key=DASHSCOPE_API_KEY,
    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
)

# 并发控制：根据 RPM=600 → QPS≈10，安全并发数设为 8
MAX_CONCURRENT_REQUESTS = 8
semaphore = threading.Semaphore(MAX_CONCURRENT_REQUESTS)


# =============================================================================
# 文档读取与转换模块
# =============================================================================

def convert_doc_to_docx(doc_path: str) -> str | None:
    """
    将 .doc 文件转换为 .docx（原地生成同名 .docx 文件）
    要求：Windows + Microsoft Word 已安装

    参数:
        doc_path (str): 原始 .doc 文件路径

    返回:
        str | None: 成功时返回新 .docx 路径，失败返回 None
    """
    try:
        import win32com.client
        docx_path = doc_path[:-4] + ".docx"
        if os.path.exists(docx_path):
            print(f"  → .docx 已存在，跳过转换: {os.path.basename(docx_path)}")
            return docx_path

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
        doc.Close()
        word.Quit()
        print(f"  → 转换成功: {os.path.basename(doc_path)} → {os.path.basename(docx_path)}")
        return docx_path
    except Exception as e:
        print(f"[ERROR] .doc 转换失败: {doc_path} | 原因: {e}")
        return None


def read_pdf(file_path: str) -> str:
    """
    使用 pdfplumber 从 PDF 文件中提取文本，保留原始布局（换行、空格等），
    有助于模型识别标题、编号和段落结构。

    参数:
        file_path (str): PDF 文件路径

    返回:
        str: 提取的全文内容，若失败则返回空字符串
    """
    try:
        full_text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # layout=True 保留原始排版，对识别 '2.7 Observing mode...' 等标题至关重要
                page_text = page.extract_text(layout=True)
                if page_text:
                    full_text += page_text + "\n"
        return full_text.strip()
    except Exception as e:
        print(f"[ERROR] 读取 PDF 失败: {file_path} | 原因: {e}")
        return ""


def read_docx(file_path: str) -> str:
    """
    从 DOCX 文件中提取纯文本内容，包括段落和表格中的文字。

    参数:
        file_path (str): DOCX 文件路径

    返回:
        str: 提取的全文内容，若失败则返回空字符串
    """
    try:
        doc = Document(file_path)
        full_text = []

        # 提取主文档段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        # 提取表格内容（逐行逐单元格）
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = " ".join(cell.text.strip().split())  # 合并多余空白
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    full_text.append(" | ".join(row_texts))  # 用分隔符连接单元格

        return "\n".join(full_text).strip()
    except Exception as e:
        print(f"[ERROR] 读取 DOCX 失败: {file_path} | 原因: {e}")
        return ""


def extract_text_from_file(file_path: str) -> str:
    """
    根据文件扩展名调用对应的读取函数。
    支持 .pdf、.docx；.doc 需先转换。

    参数:
        file_path (str): 文件路径

    返回:
        str: 提取的文本内容
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext == ".pdf":
        return read_pdf(file_path)
    elif ext == ".docx":
        return read_docx(file_path)
    else:
        print(f"[WARN] 跳过不支持的文件格式: {file_path}")
        return ""


# =============================================================================
# LLM 调用模块
# =============================================================================

def call_qwen(prompt: str) -> str:
    """
    调用 Qwen-Max 模型进行需求分解。

    注意：
    - 不使用 system 消息，避免稀释指令
    - temperature=0.1 保证输出稳定性
    - 仅返回模型生成的纯文本

    参数:
        prompt (str): 完整的提示词

    返回:
        str: 模型回复内容，若调用失败则返回错误标记
    """
    try:
        response = CLIENT.chat.completions.create(
            model="qwen-max",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            stream=False
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"[API ERROR] Qwen 调用失败: {e}")
        return "[ERROR] API 调用失败"


def call_qwen_with_semaphore(prompt: str) -> str:
    """带信号量控制的 Qwen 调用，用于并发安全"""
    with semaphore:
        return call_qwen(prompt)


# =============================================================================
# 辅助函数：按段落智能分块（增强版）
# =============================================================================

def split_document_into_chunks(text: str, max_chars: int = 50000) -> list[str]:
    """
    将长文本按自然段落切分为多个块，避免在句子中间切断。
    若单个段落超过 max_chars，则按中文句号等强制切分。

    参数:
        text (str): 原始全文
        max_chars (int): 每块最大字符数（默认 50000）

    返回:
        List[str]: 切分后的文本块列表
    """
    if len(text) <= max_chars:
        return [text] if text.strip() else []

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    current_chunk = ""

    for para in paragraphs:
        if len(para) > max_chars:
            # 超长段落：按中文句子切分
            sentences = re.split(r'(?<=[。！？；])', para)
            temp = ""
            for sent in sentences:
                if len(temp) + len(sent) <= max_chars:
                    temp += sent
                else:
                    if temp:
                        chunks.append(temp)
                    temp = sent
            if temp:
                chunks.append(temp)
            continue

        if len(current_chunk) + len(para) + 1 <= max_chars:
            current_chunk += para + "\n"
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = para + "\n"

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks


# =============================================================================
# 批次保存辅助函数
# =============================================================================

def save_batch(batch_data: list, start_idx: int, end_idx: int):
    """
    将一批需求对保存为单独的 Excel 文件。

    参数:
        batch_data (list): 当前批次的需求数据列表
        start_idx (int): 本批次起始文档序号（从1开始）
        end_idx (int): 本批次结束文档序号
    """
    if not batch_data:
        return
    df = pd.DataFrame(batch_data, columns=["高级需求", "细节需求", "来源文件"])
    filename = f"extracted_requirements_{start_idx}-{end_idx}.xlsx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    df.to_excel(output_path, index=False, engine="openpyxl")
    print(f"✅ 已保存 {filename}（共 {len(batch_data)} 条）")


# =============================================================================
# 主程序入口
# =============================================================================

def main():
    """
    主流程控制函数：
    1. 扫描输入目录，加载所有支持的文档（含 .doc 转换）
    2. 为每个文档构造提示词并调用 LLM
    3. 解析模型输出，提取有效需求对
    4. 分批次保存结果到 Excel（每 BATCH_SIZE 个有效文档一批）
    """
    # 检查输入目录是否存在
    if not os.path.exists(INPUT_DIR):
        print(f"[FATAL] 输入目录不存在: {INPUT_DIR}")
        return

    # 步骤 1: 加载所有文档（支持 .doc 自动转换）
    documents = []
    print(f"[INFO] 开始扫描目录: {INPUT_DIR}")
    for filename in os.listdir(INPUT_DIR):
        file_path = os.path.join(INPUT_DIR, filename)
        if not os.path.isfile(file_path):
            continue

        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        # === 新增：处理 .doc 文件 ===
        if ext == ".doc":
            print(f"  → 检测到 .doc 文件，尝试转换: {filename}")
            new_path = convert_doc_to_docx(file_path)
            if new_path:
                file_path = new_path
                filename = os.path.basename(new_path)
                ext = ".docx"
            else:
                print(f"  ⚠️  跳过无法转换的 .doc 文件: {filename}")
                continue

        if ext in [".pdf", ".docx"]:
            print(f"  → 正在加载: {filename}")
            content = extract_text_from_file(file_path)
            if content:
                documents.append({"filename": filename, "content": content})
            else:
                print(f"  ⚠️  跳过空内容文件: {filename}")

    if not documents:
        print("[FATAL] 未加载到任何有效文档。")
        return

    print(f"[INFO] 成功加载 {len(documents)} 份文档")

    # 步骤 2: 构造提示词模板
    prompt_template = (
        "You are a requirements analyst extracting EXPLICIT hierarchical function pairs from a software requirements document.\n\n"

        "Extract a pair ONLY if ALL the following conditions are met:\n"
        "1. The high-level item is a composite functional capability (e.g., a section title like \"Manage Users\" or \"Inventory Operations\").\n"
        "2. The sub-item represents an independent functional unit — i.e., a distinct operation, management task, view, or processing capability that a user or system can perform or access.\n"
        "3. The decomposition is presented using an EXPLICIT structural pattern, such as:\n"
        "   • Numbered subsections (e.g., 3.1, 3.2) with descriptive titles\n"
        "   • Bullet points, dashes, or indented lists (e.g., • Add user, – Edit profile)\n"
        "   • Table rows where the column clearly maps child functions to a parent function\n"
        "   • Phrases like \"including:\", \"supports the following:\", \"comprising:\", or \"the system shall provide:\" immediately followed by a list of items\n"
        "   • Hierarchical outline with indentation showing containment\n"
        "4. The sub-item is NOT:\n"
        "   - A data field, attribute, or property (e.g., \"email\", \"model number\")\n"
        "   - A state, status, or lifecycle phase (e.g., \"pending\", \"approved\")\n"
        "   - A configuration option, format, or constraint (e.g., \"CSV only\", \"max 100 characters\")\n"
        "   - An enumeration of values (e.g., \"Scraped, Broken, Lost\")\n"
        "   - A step in a workflow without standalone meaning (e.g., \"click Submit\", \"enter password\")\n\n"

        "STRICT RULES:\n"
        "- DO NOT split sentences or rephrase.\n"
        "- DO NOT infer, summarize, or create pairs from unstructured paragraphs.\n"
        "- DO NOT include any explanations, numbers, markdown, or extra text.\n"
        "- If no valid pairs exist, output NOTHING.\n\n"

        "OUTPUT FORMAT:\n"
        "- English only.\n"
        "- One pair per line: \"High-Level Function -> Sub-Function\"\n"
        "- Multiple sub-functions under the same parent should each appear on separate lines.\n\n"

        "Document content:\n{document_content}"
    )

    # 步骤 3: 逐文档处理（支持长文档分块 + 并发）
    current_batch = []      # 存储当前批次的结果
    processed_count = 0     # 成功处理的有效文档计数（从1开始）

    for idx, doc in enumerate(documents, start=1):
        filename = doc["filename"]
        content = doc["content"]

        # 分块（已改为 50000 字符）
        chunks = split_document_into_chunks(content, max_chars=50000)
        print(f"\n[进度 {idx}/{len(documents)}] 正在处理: {filename} | 切分为 {len(chunks)} 个块")

        temp_results = []  # 临时存储当前文件的所有 (高级需求, 细节需求) 对

        # === 并发处理当前文档的所有 chunks ===
        with ThreadPoolExecutor(max_workers=MAX_CONCURRENT_REQUESTS * 2) as executor:
            future_to_chunk = {
                executor.submit(call_qwen_with_semaphore, prompt_template.format(document_content=chunk)): chunk
                for chunk in chunks
            }

            for future in as_completed(future_to_chunk):
                model_response = future.result()
                for line in model_response.split("\n"):
                    line = line.strip()
                    if not line or any(line.startswith(prefix) for prefix in [
                        "注意", "注：", "[", "【", "输出", "如果", "请", "以下", "无", "未"
                    ]):
                        continue

                    if "->" in line:
                        parts = line.split("->", 1)
                        high_level = parts[0].strip()
                        low_level = parts[1].strip()
                        if high_level and low_level:
                            temp_results.append((high_level, low_level))

        # === 去重：合并同一高级需求的多个细节需求 ===
        merged = defaultdict(set)
        for high, low in temp_results:
            merged[high].add(low)

        # 如果当前文档有有效结果，才计入 processed_count
        if merged:
            processed_count += 1
            for high, low_set in merged.items():
                for low in low_set:
                    current_batch.append({
                        "高级需求": high,
                        "细节需求": low,
                        "来源文件": filename
                    })

            # 检查是否达到批次大小，自动保存
            if processed_count % BATCH_SIZE == 0:
                start = processed_count - BATCH_SIZE + 1
                end = processed_count
                save_batch(current_batch, start, end)
                current_batch = []

    # 处理最后剩余的未满批次
    if current_batch:
        start = ((processed_count - 1) // BATCH_SIZE) * BATCH_SIZE + 1
        end = processed_count
        save_batch(current_batch, start, end)

    # 最终提示
    if processed_count > 0:
        print(f"\n🎉 所有文档处理完成！结果已分批保存至 '{os.path.abspath(OUTPUT_DIR)}' 目录。")
    else:
        print("\n❌ 未提取到任何有效需求对，请检查文档内容或提示词。")


# =============================================================================
# 程序入口点
# =============================================================================

if __name__ == "__main__":
    main()